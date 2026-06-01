"""Enhanced resume processing service with text extraction and analysis."""

import PyPDF2
import docx
from pathlib import Path
from typing import Dict, Optional, List
from loguru import logger
import re
from app.core.constants import ATS_KEYWORDS

class ResumeProcessor:
    """Process and extract resume information from various formats."""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
            logger.info(f"Text extracted from PDF: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            text = ""
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
            logger.info(f"Text extracted from DOCX: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_doc(file_path: str) -> str:
        """Extract text from DOC file using python-docx."""
        return ResumeProcessor.extract_text_from_docx(file_path)
    
    @staticmethod
    def extract_contact_info(text: str) -> Dict:
        """Extract contact information from resume text."""
        contact_info = {}
        
        # Extract name (usually at the beginning)
        lines = text.split('\n')
        if lines:
            contact_info['name'] = lines[0].strip()
        
        # Extract email
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Extract phone
        phone_pattern = r'(?:\+?1[-.]?)?\(?([0-9]{3})\)?[-.]?([0-9]{3})[-.]?([0-9]{4})'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact_info['phone'] = f"({phones[0][0]}){phones[0][1]}-{phones[0][2]}"
        
        # Extract LinkedIn
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin = re.findall(linkedin_pattern, text)
        if linkedin:
            contact_info['linkedin'] = linkedin[0]
        
        return contact_info
    
    @staticmethod
    def extract_skills(text: str) -> List[str]:
        """Extract skills from resume text."""
        skills = []
        
        # Common technical skills
        technical_skills = [
            'python', 'java', 'javascript', 'typescript', 'react', 'angular', 'vue',
            'node.js', 'express', 'django', 'flask', 'fastapi', 'spring', 'sql',
            'postgresql', 'mysql', 'mongodb', 'redis', 'docker', 'kubernetes',
            'aws', 'azure', 'gcp', 'git', 'linux', 'html', 'css', 'c++', 'c#',
            'golang', 'rust', 'scala', 'r', 'matlab', 'opencv', 'tensorflow',
            'pytorch', 'scikit-learn', 'pandas', 'numpy', 'spark', 'hadoop'
        ]
        
        text_lower = text.lower()
        for skill in technical_skills:
            if skill in text_lower:
                skills.append(skill.title())
        
        # Remove duplicates while preserving order
        seen = set()
        unique_skills = []
        for skill in skills:
            if skill not in seen:
                unique_skills.append(skill)
                seen.add(skill)
        
        return unique_skills
    
    @staticmethod
    def extract_experience_years(text: str) -> int:
        """Extract total years of experience from resume text."""
        patterns = [
            r'(\d+)\s*\+?\s*years?\s+of\s+experience',
            r'experience:\s*(\d+)\s*years?',
            r'(\d+)\s*years?\s+in\s+(?:the\s+)?industry',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    return int(match.group(1))
                except (ValueError, IndexError):
                    continue
        
        return 0
    
    @staticmethod
    def calculate_ats_score(text: str, job_description: str = None) -> float:
        """Calculate ATS compatibility score (0-100)."""
        score = 50  # Base score
        
        # Check for ATS-friendly formatting
        if '\t' not in text and '|' not in text:
            score += 10  # Good: no special formatting
        
        # Check keyword density
        keywords_found = 0
        text_lower = text.lower()
        for keyword in ATS_KEYWORDS:
            if keyword.lower() in text_lower:
                keywords_found += 1
        
        score += (keywords_found / len(ATS_KEYWORDS)) * 20
        
        # Check for job description match if provided
        if job_description:
            job_keywords = set(job_description.lower().split())
            resume_words = set(text_lower.split())
            matching_words = len(job_keywords & resume_words)
            score += min((matching_words / len(job_keywords)) * 20, 20)
        
        return min(score, 100)
    
    @staticmethod
    def extract_resume_data(file_path: str, job_description: str = None) -> Dict:
        """Extract structured data from resume file."""
        try:
            # Determine file type and extract text
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.pdf':
                text = ResumeProcessor.extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                text = ResumeProcessor.extract_text_from_docx(file_path)
            elif file_ext == '.doc':
                text = ResumeProcessor.extract_text_from_doc(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_ext}")
                text = ""
            
            # Extract structured data
            data = {
                'raw_text': text,
                'contact_info': ResumeProcessor.extract_contact_info(text),
                'skills': ResumeProcessor.extract_skills(text),
                'experience_years': ResumeProcessor.extract_experience_years(text),
                'ats_score': ResumeProcessor.calculate_ats_score(text, job_description),
                'total_words': len(text.split()),
            }
            
            logger.info(f"Resume data extracted successfully. ATS Score: {data['ats_score']}")
            return data
        
        except Exception as e:
            logger.error(f"Error extracting resume data: {e}")
            return {
                'raw_text': '',
                'contact_info': {},
                'skills': [],
                'experience_years': 0,
                'ats_score': 0,
                'total_words': 0,
            }

resume_processor = ResumeProcessor()
